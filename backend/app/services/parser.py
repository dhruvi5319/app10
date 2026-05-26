"""Document parser layer — PDF (PyMuPDF), DOCX (python-docx), TXT (utf-8/latin-1)."""

import io
import logging
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import HTTPException

from app.models.errors import ErrorCode

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""

    text: str
    """Full concatenated text of the document."""

    page_chunks: list[tuple[int, str]] = field(default_factory=list)
    """List of (page_number, page_text) tuples. 1-based page numbers.
    Empty list for non-PDF documents."""

    page_count: int | None = None
    """Total number of pages (PDF only); None for TXT/DOCX."""


def parse_pdf(file_bytes: bytes) -> ParsedDocument:
    """Parse a PDF file using PyMuPDF (fitz)."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        logger.warning(f"Failed to open PDF: {exc}")
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": ErrorCode.PARSE_FAILURE,
                "message": f"Could not parse PDF file: {exc}",
            },
        )

    page_chunks: list[tuple[int, str]] = []
    total_text_parts: list[str] = []

    try:
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            page_chunks.append((page_num + 1, text))  # 1-based
            total_text_parts.append(text)
    finally:
        doc.close()

    full_text = "\n\n".join(total_text_parts)

    if not full_text.strip():
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": ErrorCode.NO_EXTRACTABLE_TEXT,
                "message": (
                    "No extractable text found in PDF. "
                    "The file may be image-only or scanned without OCR."
                ),
            },
        )

    return ParsedDocument(
        text=full_text,
        page_chunks=page_chunks,
        page_count=len(page_chunks),
    )


def parse_docx(file_bytes: bytes) -> ParsedDocument:
    """Parse a DOCX file using python-docx."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.warning(f"Failed to open DOCX: {exc}")
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": ErrorCode.PARSE_FAILURE,
                "message": f"Could not parse DOCX file: {exc}",
            },
        )

    parts: list[str] = []

    # Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Extract table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    full_text = "\n".join(parts)

    if not full_text.strip():
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": ErrorCode.NO_EXTRACTABLE_TEXT,
                "message": "No extractable text found in DOCX file.",
            },
        )

    return ParsedDocument(
        text=full_text,
        page_chunks=[],
        page_count=None,
    )


def parse_txt(file_bytes: bytes) -> ParsedDocument:
    """Parse a TXT file with UTF-8/latin-1 fallback."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
        except Exception as exc:
            raise HTTPException(
                status_code=422,
                detail={
                    "error_code": ErrorCode.PARSE_FAILURE,
                    "message": f"Could not decode text file: {exc}",
                },
            )

    if not text.strip():
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": ErrorCode.NO_EXTRACTABLE_TEXT,
                "message": "Text file is empty or contains only whitespace.",
            },
        )

    return ParsedDocument(
        text=text,
        page_chunks=[],
        page_count=None,
    )


def parse_document(file_bytes: bytes, file_type: str) -> ParsedDocument:
    """Dispatcher — routes to correct parser based on file_type."""
    if file_type == "pdf":
        return parse_pdf(file_bytes)
    elif file_type == "docx":
        return parse_docx(file_bytes)
    elif file_type == "txt":
        return parse_txt(file_bytes)
    else:
        raise HTTPException(
            status_code=422,
            detail={
                "error_code": ErrorCode.INVALID_MIME_TYPE,
                "message": f"Unsupported file type: {file_type}",
            },
        )
